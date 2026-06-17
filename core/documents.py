"""Turn an uploaded marking-scheme document into rubric inputs.

Word (.docx) -> extracted text (fed to the Architect as guidelines).
PDF / image   -> a vision ImagePart (the Architect reads it directly).
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Literal, Tuple, Union

import docx

from core.images import prep_image
from core.llm import ImagePart

IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff"}


def docx_to_text(data: bytes) -> str:
    """Flatten a .docx (paragraphs + tables) to plain text."""
    d = docx.Document(io.BytesIO(data))
    out = []
    for p in d.paragraphs:
        if p.text.strip():
            out.append(p.text.strip())
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                out.append(" | ".join(cells))
    return "\n".join(out)


def classify_upload(name: str, data: bytes) -> Tuple[Literal["text", "vision"], Union[str, ImagePart]]:
    """Route a rubric/document upload to text (docx) or a vision part (pdf/image)."""
    ext = Path(name).suffix.lower()
    if ext == ".docx":
        return "text", docx_to_text(data)
    if ext == ".txt":
        return "text", data.decode("utf-8", errors="replace")
    if ext == ".pdf":
        return "vision", ImagePart(data=data, mime_type="application/pdf")
    if ext in IMAGE_EXTS:
        return "vision", prep_image(data)
    raise ValueError(f"Unsupported document type: {ext}")
