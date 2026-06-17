"""One-off: run the engine on the real Student-1 sample (national-exc-inst)."""
import pathlib
import docx

from agents.orchestrator import StudentSheet, run_batch
from core.images import load_part

BASE = pathlib.Path("data/samples/national-exc-inst")


def docx_text(name: str) -> str:
    d = docx.Document(BASE / name)
    out = []
    for p in d.paragraphs:
        if p.text.strip():
            out.append(p.text.strip())
    for tbl in d.tables:
        for row in tbl.rows:
            out.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(out)


questions = docx_text("Questions-Functional English Pharmacy Department.docx")
rubric = docx_text("Rubrics-Functional English Pharmacy Department.docx")

guidelines = f"""You are setting up the rubric for a Functional English (Pharmacy Dept) final exam, total 50 marks.

AUTHORITATIVE EXAM QUESTIONS (this is exactly what students were asked, and what they answered):
{questions}

MARKING SCHEME / RUBRIC DOCUMENT (mark allocation, per-criterion breakdown, and language policy).
NOTE: this rubric's question wording sometimes differs from the actual questions above — when they
conflict, the ACTUAL QUESTIONS above define what was asked; use this document for the MARK
allocation, the per-criterion breakdown, and the language policy:
{rubric}

Build one rubric question per actual exam question (Q1 has 4 sub-items at 5 marks each = 20;
Q2, Q3, Q4 = 10 each). For each, give the expected model answer and the markable key points.
Apply the flexible second-language policy: 70% content/understanding, 20% application/examples,
10% language; award partial marks for conceptual understanding; minor language errors cost little."""

# pages 2..14 are answers (page 1 is the cover/instructions)
pages = sorted((BASE / "student 1-answer sheet").glob("*.jpg"),
               key=lambda p: int(p.stem.split("_")[-1]))
answer_pages = [p for p in pages if int(p.stem.split("_")[-1]) >= 2]
print(f"Loading {len(answer_pages)} answer pages...")
parts = [load_part(p) for p in answer_pages]

view = run_batch("national-exc", "Functional English — Pharmacy",
                 sheets=[StudentSheet("student_1", parts)], guidelines=guidelines)

print("\n===== RUBRIC =====")
print(view.rubric.model_dump_json(indent=2))
r = view.results["student_1"]
print(f"\n===== STUDENT_1: {r.total_awarded}/{r.total_max} ({len(r.open_flags)} flagged) =====")
print(r.model_dump_json(indent=2))
print("\n===== SUMMARY =====")
print(view.summaries["student_1"].model_dump_json(indent=2))
